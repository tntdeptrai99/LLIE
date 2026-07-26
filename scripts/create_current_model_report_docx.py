from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BENCH = ROOT / "reports" / "benchmarks" / "current_model_20260726"
PIPE = ROOT / "reports" / "benchmarks" / "current_pipeline_20260726"
FIG = ROOT / "reports" / "figures" / "current_pipeline_20260726" / "real_camera_visual"
OUT = BENCH / "bao_cao_benchmark_model_hien_tai_current_model.docx"


def shade(cell, color="F2F4F7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def text(cell, value, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(9)


def borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D9E2F3")


def finish_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    borders(table)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                shade(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, h in zip(table.rows[0].cells, headers):
        text(cell, h, True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            text(cell, value)
    finish_table(table)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_image(doc, path: Path, caption: str, width=6.2):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)


def load_csv(path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles(doc)

    p = doc.add_paragraph()
    r = p.add_run("Báo cáo benchmark model hiện tại")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    doc.add_paragraph("Ghost-ESP current deployed ONNX | Current model 2026-07-26")

    doc.add_heading("1. Phạm vi benchmark", level=1)
    add_bullets(
        doc,
        [
            "Benchmark chất lượng trên train/val dataset bằng PC ONNX.",
            "Đối chiếu log train hiện có của model current.",
            "So sánh PC ONNX với board output qua UART tensor dump.",
            "So sánh thị giác qua board: input/preprocess và board AI output.",
            "Thống kê chi phí inference/pipeline trên board từ UART status log.",
        ],
    )

    doc.add_heading("2. Thông số model", level=1)
    add_table(
        doc,
        ["Hạng mục", "Giá trị"],
        [
            ["Kiến trúc", "Ghost-ESP current deployed ONNX"],
            ["Input", "input_rgb, 1x3x96x96, float"],
            ["Output", "enhanced_rgb, 1x3x96x96, uint8"],
            ["Số tham số", "5,130"],
            ["Model size INT8/ONNX", "70.79 KB"],
            ["Layout khớp board", "raw_nchw"],
        ],
    )

    doc.add_heading("3. Train log hiện có", level=1)
    add_table(
        doc,
        ["Epoch", "Train loss", "Val loss", "Val PSNR", "Val SSIM", "Teacher PSNR", "Teacher SSIM"],
        [["80", "0.288501", "0.109468", "17.0453", "0.8193", "25.3867", "0.9536"]],
    )
    doc.add_paragraph("Ghi chú: đây là log train cũ còn trong project, chưa phải retrain mới từ cùng seed/config.")

    doc.add_heading("4. PC ONNX benchmark theo dataset split", level=1)
    summary = json.loads((BENCH / "pc_dataset_split_summary.json").read_text(encoding="utf-8"))
    rows = []
    for dataset, s in summary.items():
        rows.append(
            [
                dataset,
                s["n"],
                f"{s['psnr_mean']:.4f}",
                f"{s['ssim_mean']:.4f}",
                f"{s['mae_mean']:.6f}",
                f"{s['inference_ms_pc_mean']:.4f}",
            ]
        )
    add_table(doc, ["Dataset", "N", "PSNR", "SSIM", "MAE", "PC ms/frame"], rows)

    doc.add_heading("5. PC-board equivalence", level=1)
    add_table(
        doc,
        ["Hạng mục", "Kết quả"],
        [
            ["Best layout", "raw_nchw"],
            ["Exact match", "98.09%"],
            ["MAE lượng tử", "0.022 / 255"],
            ["Max abs", "3"],
            ["Cosine", "1.000000"],
        ],
    )
    doc.add_paragraph("Kết luận: board AI output khớp PC ONNX rất sát với dump hiện tại.")

    doc.add_heading("6. Thị giác qua board", level=1)
    add_image(doc, FIG / "board_input_vs_ai_output_contact_x4.png", "Input/preprocess so với board AI output.")
    metrics = load_csv(FIG / "image_metrics.csv")
    metric_rows = []
    by_name = {r["name"]: r for r in metrics}
    for key in [
        "brightness_mean",
        "contrast_std",
        "contrast_p01_p99",
        "saturation_mean",
        "sharpness_laplacian_abs_mean",
        "clip_0_ratio_rgb",
        "clip_255_ratio_rgb",
    ]:
        metric_rows.append([key, f"{float(by_name['input_preprocess'][key]):.6f}", f"{float(by_name['ai_output'][key]):.6f}"])
    add_table(doc, ["Metric", "Input/preprocess", "Board AI output"], metric_rows)
    add_image(doc, FIG / "luma_histogram_input_vs_ai_output.png", "Histogram luminance input so với AI output.")

    doc.add_heading("7. Chi phí inference/pipeline trên board", level=1)
    timing = load_csv(BENCH / "board_timing_summary.csv")
    names = {
        "camera_display_ms": "Camera/LCD display ms",
        "preprocess_ms": "Preprocess ms",
        "inference_ms": "Inference ms",
        "postprocess_ms": "Postprocess ms",
        "total_ms": "Total pipeline ms",
        "fps": "FPS",
    }
    timing_rows = [
        [names[r["metric"]], r["n"], r["min"], r["mean"], r["max"], r["p95"]]
        for r in timing
        if r["metric"] in names
    ]
    add_table(doc, ["Metric", "N", "Min", "Mean", "Max", "P95"], timing_rows)
    add_bullets(
        doc,
        [
            "AI inference trên board là nút cổ chai chính: mean khoảng 170.3 ms/frame.",
            "Tổng pipeline mean khoảng 191.6 ms/frame, tương đương khoảng 5 FPS.",
            "Preprocess và postprocess nhỏ hơn nhiều so với inference.",
        ],
    )

    doc.add_heading("8. Kết luận", level=1)
    add_bullets(
        doc,
        [
            "PC-board output khớp tốt, nên cần tách lỗi model quality khỏi lỗi LCD/render buffer.",
            "Model làm sáng và tăng contrast ảnh board, nhưng saturation giảm và có clipping trắng nhẹ.",
            "Board inference hiện là chi phí lớn nhất, không phải preprocess/postprocess.",
            "Cần log profiling mới dạng prof/p cycles nếu muốn báo cáo chi phí board chi tiết hơn theo từng stage CPU cycle.",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
