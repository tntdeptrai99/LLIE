from __future__ import annotations

import csv
import json
import statistics
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
DATE = "20260726"
NEW_DATE = "20260807"

OUT_DIR = ROOT / "reports" / "benchmarks" / f"final_comprehensive_{NEW_DATE}"
FIG_DIR = ROOT / "reports" / "figures" / f"final_comprehensive_{NEW_DATE}"
OUT_DOCX = OUT_DIR / f"bao_cao_hoan_chinh_darkghost_espnet_train_pc_board_{NEW_DATE}_v2.docx"
OUT_MD = OUT_DIR / f"bao_cao_hoan_chinh_darkghost_espnet_train_pc_board_{NEW_DATE}_v2.md"

METRICS = ROOT / "reports" / "metrics"
CURRENT_MODEL = ROOT / "reports" / "benchmarks" / f"current_model_{DATE}"
CURRENT_PIPE = ROOT / "reports" / "benchmarks" / f"current_pipeline_{DATE}"
CURRENT_REAL_FIG = ROOT / "reports" / "figures" / f"current_pipeline_{DATE}" / "real_camera_visual"
METHOD_FIG = ROOT / "reports" / "figures" / f"darkghost_espnet_{DATE}" / "darkghost_espnet_training_deployment_framework.png"
PAPER_FIG = ROOT / "reports" / "figures" / f"paper_style_{NEW_DATE}"
COMP_FIG = ROOT / "reports" / "figures" / f"component_ablation_{DATE}" / "component_ablation_effect.png"
ARCH_FIG = ROOT / "reports" / "figures" / f"architecture_ablation_{DATE}" / "architecture_cost_comparison.png"
BEST_VIS_DIR = ROOT / "reports" / "figures" / "optuna_best_300e"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def parse_summary(path: Path) -> dict[str, str]:
    out: dict[str, str] = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        if "=" in line:
            key, value = line.split("=", 1)
            out[key.strip()] = value.strip()
    return out


def fnum(value: str | float | int, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def kib(value: int | float) -> str:
    return f"{float(value) / 1024:.2f} KiB"


def selected_research_rows() -> list[dict[str, str]]:
    names = [
        ("MAE Baseline 300e", "mae_baseline_300e"),
        ("Base Model 300e", "base_model_300e"),
        ("No Dark-Map 300e", "no_dark_300e"),
        ("No Teacher KD 300e", "no_teacher_300e"),
        ("Optuna Best 300e", "optuna_best_300e"),
    ]
    rows: list[dict[str, str]] = []
    for label, slug in names:
        paths = sorted(METRICS.glob(f"lol_test_{slug}_{slug}_summary.txt"))
        if not paths:
            paths = sorted(METRICS.glob(f"*{slug}*summary.txt"))
        if not paths:
            continue
        data = parse_summary(paths[0])
        rows.append(
            {
                "label": label,
                "artifact": slug,
                "n": data.get("count", ""),
                "psnr": data.get("psnr", ""),
                "ssim": data.get("ssim", ""),
                "source": str(paths[0]),
            }
        )
    return rows


def qdq_summary() -> dict[str, str]:
    rows = read_csv(METRICS / "qdq_drift_plateau_score_best_monitor.csv")
    return {
        "n": str(len(rows)),
        "mean_abs": fnum(statistics.mean(float(r["mean_abs"]) for r in rows), 6),
        "max_abs": fnum(max(float(r["max_abs"]) for r in rows), 6),
        "rmse": fnum(statistics.mean(float(r["rmse"]) for r in rows), 6),
        "psnr": fnum(statistics.mean(float(r["psnr_fp32_vs_qdq"]) for r in rows), 4),
        "psnr_min": fnum(min(float(r["psnr_fp32_vs_qdq"]) for r in rows), 4),
    }


def make_best_contact_sheet() -> Path | None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    imgs = sorted(BEST_VIS_DIR.glob("*.png"))[:6]
    if not imgs:
        return None
    loaded = [Image.open(p).convert("RGB") for p in imgs]
    thumb_w = 330
    thumbs = []
    for im in loaded:
        ratio = thumb_w / im.width
        thumbs.append(im.resize((thumb_w, int(im.height * ratio)), Image.Resampling.LANCZOS))
    label_h = 28
    pad = 18
    cols = 2
    rows = (len(thumbs) + cols - 1) // cols
    cell_h = max(t.height for t in thumbs) + label_h + pad
    canvas = Image.new("RGB", (cols * thumb_w + (cols + 1) * pad, rows * cell_h + pad), "white")
    draw = ImageDraw.Draw(canvas)
    for idx, thumb in enumerate(thumbs):
        x = pad + (idx % cols) * (thumb_w + pad)
        y = pad + (idx // cols) * cell_h
        draw.text((x, y), f"Mẫu LOL test {idx:02d}: Input - Output - Ground Truth", fill=(20, 20, 20))
        canvas.paste(thumb, (x, y + label_h))
    out = FIG_DIR / "best_research_checkpoint_contact_sheet.png"
    canvas.save(out, quality=95)
    return out


def set_cell(cell, text: str, bold: bool = False, center: bool = False, size: float = 8.6) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "CDD6E1")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_borders(table)
    for i, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
        shade(cell, "E8EEF5")
        set_cell(cell, header, bold=True, center=True, size=8.4)
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, (cell, value) in enumerate(zip(cells, row)):
            set_cell(cell, value, center=len(str(value)) <= 14, size=8.2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                cell.width = Inches(widths[i])
    doc.add_paragraph()


def add_figure(doc: Document, path: Path | None, caption: str, width: float = 6.3) -> None:
    if path is None or not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_note(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph()


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DarkGhost-ESPNet")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(23, 54, 93)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Báo cáo tổng hợp train - PC - board cho mô hình LLIE trên STM32")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(80, 80, 80)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Ngày tạo: 07/08/2026 | Artifact root: {ROOT}")
    doc.add_paragraph()


def collect() -> dict:
    return {
        "research": selected_research_rows(),
        "pc": json.loads((CURRENT_MODEL / "pc_dataset_split_summary.json").read_text(encoding="utf-8")),
        "timing": {r["metric"]: r for r in read_csv(CURRENT_MODEL / "board_timing_summary.csv")},
        "visual": {r["name"]: r for r in read_csv(CURRENT_REAL_FIG / "image_metrics.csv")},
        "arch_cost": read_csv(ROOT / "reports" / "benchmarks" / f"architecture_ablation_{DATE}" / "architecture_ablation_metrics.csv"),
        "arch_train": read_csv(ROOT / "reports" / "benchmarks" / f"architecture_ablation_saturation_{DATE}" / "best_available_saturation_summary.csv"),
        "component": read_csv(ROOT / "reports" / "benchmarks" / f"component_ablation_{DATE}" / "component_ablation_summary.csv"),
        "qdq": qdq_summary(),
    }


def write_markdown(data: dict, best_contact: Path | None) -> None:
    lines = [
        "# DarkGhost-ESPNet: báo cáo tổng hợp train - PC - board",
        "",
        "Báo cáo này chỉ dùng artifact có sẵn trong project tại thời điểm 07/08/2026 (300 epochs).",
        "",
        "## Kết quả chính",
        "",
        "- Best research checkpoint: `optuna_best_300e` (cập nhật sau khi đánh giá).",
        "- Current deployed ONNX: `stm32/onnx/ghost_esp_dark_w12_m24_enhanced_rgb_u8out_tail_simplified_nchw.onnx`.",
        "- PC-board equivalence: layout `raw_nchw`, exact 98.09%, MAE lượng tử 0.022/255, cosine 1.000000.",
        "- Board pipeline hiện tại: inference mean 170.3256 ms, total mean 191.5581 ms, FPS mean 4.9070.",
        "",
        "## Traceability",
        "",
        f"- Báo cáo Word: `{OUT_DOCX}`",
        f"- Contact sheet best checkpoint: `{best_contact}`",
    ]
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def create_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    data = collect()
    best_contact = make_best_contact_sheet()
    write_markdown(data, best_contact)

    doc = Document()
    style_doc(doc)
    add_title(doc)

    add_note(
        doc,
        "Tóm tắt xác thực",
        "Báo cáo tách rõ ba lớp artifact: checkpoint nghiên cứu tốt nhất đã khôi phục, ONNX đang deploy trên firmware, và kết quả đo thực tế trên board. Vì vậy các claim train/PC/board không bị trộn lẫn giữa các model khác nhau.",
    )

    doc.add_heading("1. Phạm vi và nguồn số liệu", level=1)
    doc.add_paragraph(
        "Các số liệu được lấy từ log/checkpoint/CSV/JSON/PNG đang tồn tại trong project sau khi khôi phục. "
        "Không dùng số liệu ước lượng để thay thế kết quả benchmark. Các benchmark ngắn hoặc bị ngắt được ghi rõ trạng thái."
    )
    add_table(
        doc,
        ["Nhóm dữ liệu", "Artifact nguồn", "Vai trò trong báo cáo"],
        [
            ["Train/eval research", "reports/metrics/*.txt, reports/figures/ghost_esp_*", "Chọn checkpoint tốt nhất và hình so sánh trực quan"],
            ["PC ONNX", "reports/benchmarks/current_model_20260726/pc_dataset_split_summary.json", "Đánh giá ONNX đang deploy trên các dataset split"],
            ["QDQ/quantization", "reports/metrics/qdq_drift_plateau_score_best_monitor.csv", "Đo sai lệch FP32 - QDQ"],
            ["Board", "board_timing_summary.csv, pc_board_equivalence.log, real_camera_visual", "Đo latency, tương đương PC-board và chất lượng thị giác thực tế"],
        ],
        widths=[1.5, 2.7, 2.4],
    )

    doc.add_heading("2. Mô hình đề xuất và pipeline train-deploy", level=1)
    doc.add_paragraph(
        "Tên mô hình dùng trong báo cáo: DarkGhost-ESPNet. Kiến trúc sử dụng student Ghost/ESP nhẹ ở kích thước 96x96, "
        "hướng dẫn vùng tối bằng dark map, teacher RGB theo hướng Retinex/VD, adaptive loss theo dark map và tối ưu trọng số loss bằng Optuna. "
        "Khi deploy lên STM32 chỉ dùng student đã lượng tử hóa, không deploy teacher."
    )
    add_figure(doc, METHOD_FIG, "Hình 1. Sơ đồ train - distillation - deploy của DarkGhost-ESPNet.", width=6.5)

    doc.add_heading("3. Best research checkpoint đã khôi phục", level=1)
    research_rows = [
        [r["label"], r["n"], fnum(r["psnr"]), fnum(r["ssim"]), r["artifact"]]
        for r in data["research"]
    ]
    add_table(doc, ["Checkpoint/pha train", "N", "PSNR", "SSIM", "Artifact"], research_rows, widths=[1.55, 0.35, 0.55, 0.55, 3.5])
    add_note(
        doc,
        "Kết luận",
        "Checkpoint tốt nhất được chạy lại lên 300 epochs là Optuna Best 300e. Đây là số nên dùng làm kết quả research chính thống để đảm bảo tính đồng bộ công bằng.",
    )
    add_figure(doc, best_contact, "Hình 2. Contact sheet các mẫu đánh giá của best research checkpoint.", width=6.5)

    doc.add_heading("4. Ablation thành phần: Optuna, dark-map loss, teacher", level=1)
    comp_rows = []
    for r in data["component"]:
        comp_rows.append([r["variant"], r["count"], r["epoch"], fnum(r["psnr"]), fnum(r["ssim"]), r["notes"]])
    add_table(doc, ["Biến thể", "N", "Epoch/source", "PSNR", "SSIM", "Ghi chú"], comp_rows, widths=[1.35, 0.3, 0.72, 0.55, 0.55, 3.0])
    add_figure(doc, COMP_FIG, "Hình 3. Ảnh hưởng của Optuna, adaptive dark-map loss và teacher KD.", width=6.2)
    add_note(
        doc,
        "Diễn giải đúng",
        "Hai dòng 'No adaptive dark-map loss' và 'No teacher KD' là ablation CUDA 30 epoch để nhận diện đóng góp thành phần. Chúng không cùng budget với best checkpoint plateau/refine, nên dùng như bằng chứng xu hướng chứ không tuyên bố là so sánh hội tụ tuyệt đối.",
    )

    doc.add_heading("5. Optuna và đường cong train/validation", level=1)
    doc.add_paragraph(
        "Optuna thực tế trong project có 20 complete trials. Các figure bên dưới dùng log thật hiện có để minh họa quá trình tối ưu và đường cong train/validation; không nội suy thành 500 trials."
    )
    add_figure(doc, PAPER_FIG / "optuna_optimization_progress.png", "Hình 4. Tiến trình Optuna theo trial thực tế.", width=6.1)
    add_figure(doc, PAPER_FIG / "optuna_best_loss_weights.png", "Hình 5. Bộ trọng số loss tốt nhất từ Optuna.", width=6.1)
    add_figure(doc, PAPER_FIG / "training_validation_curves_project.png", "Hình 6. Đường cong train/validation từ log project.", width=6.2)
    add_figure(doc, PAPER_FIG / "component_ablation_training_curves.png", "Hình 7. Đường cong train của các ablation thành phần.", width=6.2)

    doc.add_heading("6. Benchmark kiến trúc mô hình", level=1)
    cost_rows = []
    for r in data["arch_cost"]:
        cost_rows.append(
            [
                r["architecture"],
                r["params"],
                r["model_size_fp32_kib"],
                r["model_size_int8_est_kib"],
                r["macs"],
                fnum(r["latency_ms_mean"]),
                fnum(r["psnr"]) if r["psnr"] else "-",
                fnum(r["ssim"]) if r["ssim"] else "-",
                r["note"],
            ]
        )
    add_table(
        doc,
        ["Architecture", "Params", "FP32 KiB", "INT8 est KiB", "MACs", "PC ms", "PSNR", "SSIM", "Ghi chú"],
        cost_rows,
        widths=[0.9, 0.55, 0.55, 0.65, 0.85, 0.45, 0.45, 0.45, 2.3],
    )
    add_figure(doc, ARCH_FIG, "Hình 8. So sánh chi phí kiến trúc theo tham số, MACs và latency PC.", width=6.2)
    train_rows = [
        [r["architecture"], r["status"], r["epochs_ran"], r["best_epoch"], fnum(r["best_psnr"]), fnum(r["best_ssim"]), fnum(r["best_mae"], 6)]
        for r in data["arch_train"]
    ]
    add_table(doc, ["Architecture", "Trạng thái", "Epoch chạy", "Best epoch", "PSNR", "SSIM", "MAE"], train_rows, widths=[1.0, 1.35, 0.65, 0.65, 0.55, 0.55, 0.65])

    doc.add_heading("7. Current deployed ONNX trên PC", level=1)
    pc_rows = []
    for name, r in data["pc"].items():
        pc_rows.append([name, str(r["n"]), fnum(r["psnr_mean"]), fnum(r["ssim_mean"]), fnum(r["mae_mean"], 6), fnum(r["inference_ms_pc_mean"])])
    add_table(doc, ["Dataset split", "N", "PSNR", "SSIM", "MAE", "PC ONNX ms/frame"], pc_rows, widths=[1.4, 0.45, 0.55, 0.55, 0.65, 0.9])
    add_note(
        doc,
        "Lưu ý quan trọng",
        "ONNX đang deploy đạt tốt nhất trên LOL_v2_Real_val (PSNR 19.4780, SSIM 0.7858). Kết quả LOL_train/LOL_val thấp hơn do khác distribution/artifact deploy, nên không dùng để phủ nhận best research checkpoint mà dùng để đánh giá model deploy hiện tại.",
    )

    doc.add_heading("8. Quantization và tương đương PC-board", level=1)
    q = data["qdq"]
    add_table(
        doc,
        ["N ảnh", "Mean abs", "Max abs lớn nhất", "RMSE mean", "PSNR FP32-QDQ mean", "PSNR min"],
        [[q["n"], q["mean_abs"], q["max_abs"], q["rmse"], q["psnr"], q["psnr_min"]]],
        widths=[0.55, 0.85, 1.0, 0.85, 1.2, 0.8],
    )
    add_table(
        doc,
        ["Kiểm chứng", "Kết quả"],
        [
            ["Input/output ONNX", "input_rgb -> enhanced_rgb_QuantizeLinear_Output, uint8, 1x3x96x96"],
            ["Layout board tốt nhất", "raw_nchw"],
            ["Exact match lượng tử", "98.09%"],
            ["MAE lượng tử", "0.022 / 255"],
            ["Max abs lượng tử", "3 / 255"],
            ["Cosine", "1.000000"],
        ],
        widths=[1.7, 4.4],
    )

    doc.add_heading("9. Benchmark board: runtime và tài nguyên", level=1)
    t = data["timing"]
    board_rows = [
        ["Camera/LCD display", fnum(t["camera_display_ms"]["mean"]), fnum(t["camera_display_ms"]["min"]), fnum(t["camera_display_ms"]["max"])],
        ["Preprocess", fnum(t["preprocess_ms"]["mean"]), fnum(t["preprocess_ms"]["min"]), fnum(t["preprocess_ms"]["max"])],
        ["Inference", fnum(t["inference_ms"]["mean"]), fnum(t["inference_ms"]["min"]), fnum(t["inference_ms"]["max"])],
        ["Postprocess", fnum(t["postprocess_ms"]["mean"]), fnum(t["postprocess_ms"]["min"]), fnum(t["postprocess_ms"]["max"])],
        ["Total pipeline", fnum(t["total_ms"]["mean"]), fnum(t["total_ms"]["min"]), fnum(t["total_ms"]["max"])],
        ["FPS", fnum(t["fps"]["mean"]), fnum(t["fps"]["min"]), fnum(t["fps"]["max"])],
    ]
    add_table(doc, ["Stage", "Mean", "Min", "Max"], board_rows, widths=[1.7, 0.8, 0.8, 0.8])
    add_table(
        doc,
        ["Artifact", "Số liệu"],
        [
            ["Best research QDQ Cube.AI", "Flash 55,464 B; weights 5,944 B; total RAM 441,820 B; activations 425,260 B"],
            ["Current deployed ONNX", "70.79 KiB; khoảng 5,130 params; input/output uint8 NCHW 96x96"],
            ["Current firmware report", "MACC khoảng 19.386M; weights 5.79 KiB; activations khoảng 518.96 KiB"],
        ],
        widths=[1.7, 4.4],
    )

    doc.add_heading("10. Benchmark thị giác thực tế trên board", level=1)
    v = data["visual"]
    visual_rows = []
    for name in ["input_preprocess", "ai_output"]:
        r = v[name]
        visual_rows.append(
            [
                name,
                fnum(r["brightness_mean"]),
                fnum(r["contrast_std"]),
                fnum(r["saturation_mean"]),
                fnum(r["sharpness_laplacian_abs_mean"]),
                fnum(r["clip_0_ratio_rgb"], 6),
                fnum(r["clip_255_ratio_rgb"], 6),
            ]
        )
    add_table(
        doc,
        ["Ảnh", "Brightness", "Contrast", "Saturation", "Sharpness", "Clip 0", "Clip 255"],
        visual_rows,
        widths=[0.8, 0.8, 0.8, 0.8, 0.85, 0.65, 0.65],
    )
    add_figure(doc, CURRENT_REAL_FIG / "board_input_vs_ai_output_contact_x4.png", "Hình 9. So sánh thị giác thực tế trên board: input preprocess và AI output.", width=6.2)
    add_figure(doc, CURRENT_REAL_FIG / "luma_histogram_input_vs_ai_output.png", "Hình 10. Histogram độ sáng trước/sau trên frame board thực tế.", width=6.2)

    doc.add_heading("11. Kết luận và cách dùng trong paper", level=1)
    doc.add_paragraph(
        "Kết quả mạnh nhất của project hiện tại là chuỗi bằng chứng: DarkGhost-ESPNet đạt PSNR/SSIM tốt trên checkpoint nghiên cứu đã khôi phục; "
        "ONNX deploy hiện tại chạy được trên PC với dataset split và tương đương rất sát với board dump; board đo được latency end-to-end và có visual metric thực tế. "
        "Khi viết paper, nên báo cáo best research checkpoint cho phần chất lượng mô hình, còn phần embedded nên báo cáo current deployed ONNX và board timing như một artifact triển khai riêng."
    )
    add_note(
        doc,
        "Giới hạn cần ghi rõ",
        "Các architecture baseline chưa được train đồng nhất đến bão hòa đầy đủ. Một số baseline là cost-only hoặc controlled/pilot run. Optuna hiện có 20 trials thật, không phải 500 trials. Đây là cách trình bày trung thực nhất với artifact hiện có.",
    )

    doc.save(OUT_DOCX)


def structural_check() -> dict[str, int]:
    with zipfile.ZipFile(OUT_DOCX) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
    d = Document(OUT_DOCX)
    return {
        "paragraphs": len(d.paragraphs),
        "tables": len(d.tables),
        "images": len(media),
    }


if __name__ == "__main__":
    create_docx()
    print(json.dumps({"docx": str(OUT_DOCX), "md": str(OUT_MD), "fig_dir": str(FIG_DIR), "check": structural_check()}, ensure_ascii=False, indent=2))
