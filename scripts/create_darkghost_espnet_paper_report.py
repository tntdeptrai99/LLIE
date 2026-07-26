from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATE = "20260726"

BENCH_DARK = ROOT / "reports" / "benchmarks" / f"darkghost_espnet_{DATE}"
FIG_DARK = ROOT / "reports" / "figures" / f"darkghost_espnet_{DATE}"
BENCH_MODEL = ROOT / "reports" / "benchmarks" / f"current_model_{DATE}"
BENCH_PIPE = ROOT / "reports" / "benchmarks" / f"current_pipeline_{DATE}"
FIG_MODEL = ROOT / "reports" / "figures" / f"current_model_{DATE}"
FIG_REAL = ROOT / "reports" / "figures" / f"current_pipeline_{DATE}" / "real_camera_visual"
BENCH_ARCH = ROOT / "reports" / "benchmarks" / f"architecture_ablation_{DATE}"
FIG_ARCH = ROOT / "reports" / "figures" / f"architecture_ablation_{DATE}"
BENCH_ARCH_TRAIN = ROOT / "reports" / "benchmarks" / f"architecture_ablation_unified_train_{DATE}"
FIG_ARCH_TRAIN = ROOT / "reports" / "figures" / f"architecture_ablation_unified_train_{DATE}"

OUT_MD = BENCH_DARK / "darkghost_espnet_miwai_style_report.md"
OUT_DOCX = BENCH_DARK / "bao_cao_darkghost_espnet_miwai_style.docx"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def shade(cell, color: str = "EAF2F8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_text(cell, value, bold: bool = False, align_center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(8.8)


def set_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C8D6E5")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_borders(table)
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell)
        set_cell_text(cell, header, bold=True, align_center=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value, align_center=len(str(value)) < 16)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.5) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Thiếu hình: {path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("555555")


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 15, "17365D"),
        ("Heading 2", 12, "1F4E79"),
        ("Heading 3", 11, "2F5597"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def fmt(v, digits=4) -> str:
    return f"{float(v):.{digits}f}"


def collect_data() -> dict:
    top = read_csv(BENCH_DARK / "train_log_summary_top10.csv")
    best = top[0]
    pc = read_json(BENCH_MODEL / "pc_dataset_split_summary.json")
    board_timing_rows = read_csv(BENCH_MODEL / "board_timing_summary.csv")
    board_timing = {row["metric"]: row for row in board_timing_rows}
    visual = read_csv(FIG_REAL / "image_metrics.csv")
    visual_by_name = {row["name"]: row for row in visual}
    arch_rows = read_csv(BENCH_ARCH / "architecture_ablation_metrics.csv") if (BENCH_ARCH / "architecture_ablation_metrics.csv").exists() else []
    arch_train_rows = (
        read_csv(BENCH_ARCH_TRAIN / "unified_architecture_train_summary.csv")
        if (BENCH_ARCH_TRAIN / "unified_architecture_train_summary.csv").exists()
        else []
    )
    equiv_log = (BENCH_PIPE / "pc_board_equivalence.log").read_text(encoding="utf-8", errors="replace")
    return {
        "best": best,
        "top": top,
        "pc": pc,
        "board_timing": board_timing,
        "visual": visual_by_name,
        "arch": arch_rows,
        "arch_train": arch_train_rows,
        "equiv_log": equiv_log,
    }


def write_markdown(data: dict) -> None:
    best = data["best"]
    board = data["board_timing"]
    pc = data["pc"]
    visual = data["visual"]
    arch = data["arch"]
    arch_train = data["arch_train"]
    framework = FIG_DARK / "darkghost_espnet_training_deployment_framework.png"
    serious = FIG_MODEL / "serious_current_method_ghost_esp_darkmap_retinex_optuna.png"
    contact = FIG_REAL / "board_input_vs_ai_output_contact_x4.png"
    hist = FIG_REAL / "luma_histogram_input_vs_ai_output.png"
    arch_fig = FIG_ARCH / "architecture_cost_comparison.png"
    arch_train_fig = FIG_ARCH_TRAIN / "unified_architecture_quality_comparison.png"

    pc_lines = [
        "| Dataset split | N | PSNR | SSIM | MAE | PC ONNX ms/frame |",
        "|---|---:|---:|---:|---:|---:|",
    ]
    for name, row in pc.items():
        pc_lines.append(
            f"| {name} | {row['n']} | {fmt(row['psnr_mean'])} | {fmt(row['ssim_mean'])} | "
            f"{fmt(row['mae_mean'], 6)} | {fmt(row['inference_ms_pc_mean'])} |"
        )

    visual_lines = [
        "| Metric | Input/preprocess | Board AI output |",
        "|---|---:|---:|",
    ]
    for key in [
        "brightness_mean",
        "contrast_std",
        "saturation_mean",
        "sharpness_laplacian_abs_mean",
        "clip_0_ratio_rgb",
        "clip_255_ratio_rgb",
    ]:
        visual_lines.append(
            f"| {key} | {fmt(visual['input_preprocess'][key], 6)} | {fmt(visual['ai_output'][key], 6)} |"
        )

    arch_lines = [
        "| Architecture | Params | FP32 KiB | INT8 est. KiB | MACs | CPU ms | PSNR | SSIM | MAE | Source |",
        "|---|---:|---:|---:|---:|---:|---:|---:|---:|---|",
    ]
    for row in arch:
        arch_lines.append(
            f"| {row['architecture']} | {int(row['params']):,} | {row['model_size_fp32_kib']} | "
            f"{row['model_size_int8_est_kib']} | {int(row['macs']):,} | {row['latency_ms_mean']} | "
            f"{row['psnr'] or '-'} | {row['ssim'] or '-'} | {row['mae'] or '-'} | {row['quality_source']} |"
        )

    arch_train_lines = [
        "| Architecture | Params | MACs | Best epoch | PSNR | SSIM | MAE |",
        "|---|---:|---:|---:|---:|---:|---:|",
    ]
    for row in arch_train:
        arch_train_lines.append(
            f"| {row['architecture']} | {int(row['params']):,} | {int(row['macs']):,} | "
            f"{row['best_epoch']} | {fmt(row['best_psnr'])} | {fmt(row['best_ssim'])} | {fmt(row['best_mae'], 6)} |"
        )

    text = f"""# DarkGhost-ESPNet: báo cáo benchmark kiểu MIWAI

**Ngày tạo:** 2026-07-26  
**Tên mô hình đề xuất:** DarkGhost-ESPNet  
**Biến thể triển khai:** DarkGhost-ESPNet-Tiny96  
**ONNX đang triển khai:** `stm32/onnx/ghost_esp_dark_w12_m24_enhanced_rgb_u8out_tail_simplified_nchw.onnx`

## Tóm tắt

Báo cáo này đồng bộ các số liệu mới nhất đang có trong project cho hướng mô hình Ghost-ESP có hướng dẫn dark-map, teacher Retinex/VD trong quá trình huấn luyện, loss thích nghi theo vùng tối và tối ưu siêu tham số bằng Optuna. Điểm cần tách rõ: trên STM32 chỉ triển khai student `DarkGhost-ESPNet-Tiny96`; teacher, loss và Optuna là cơ chế huấn luyện, không chạy trong firmware.

![Framework]({framework.as_posix()})

## Phương pháp đề xuất

DarkGhost-ESPNet dùng student nhẹ dựa trên Ghost-ESP để tăng cường ảnh thiếu sáng 96x96. Dark-map được tạo từ ảnh đầu vào để nhấn mạnh các vùng thiếu sáng, sau đó điều biến loss theo vùng tối. Teacher RGB dựa trên Retinex/VD cung cấp tín hiệu tham chiếu trong huấn luyện/distillation. Optuna được dùng để tìm cấu hình trọng số loss và tham số huấn luyện tốt hơn.

![Serious method]({serious.as_posix()})

## Kết quả train mới nhất và đồng bộ

Run tốt nhất theo log hiện có:

- Experiment: `{best['experiment']}`
- Số epoch: `{best['epochs']}`
- Best PSNR: **{fmt(best['best_psnr'])} dB** tại epoch `{best['best_psnr_epoch']}`
- Best SSIM: **{fmt(best['best_ssim'])}** tại epoch `{best['best_ssim_epoch']}`
- Epoch cuối: `{best['last_epoch']}`, val PSNR `{fmt(best['last_val_psnr'])}`, val SSIM `{fmt(best['last_val_ssim'])}`
- Teacher/reference log: PSNR `{fmt(best['teacher_psnr'])}`, SSIM `{fmt(best['teacher_ssim'])}`

Ghi chú: đây là số liệu đồng bộ từ log mới nhất trong project, không còn dùng riêng số epoch 80 cũ làm đại diện chính.

## Benchmark PC ONNX theo dataset split

{chr(10).join(pc_lines)}

## Benchmark kiến trúc model

Benchmark này so sánh ablation kiến trúc tương tự paper trên input 96x96. Chi phí được đo trực tiếp bằng PyTorch CPU; PSNR/SSIM/MAE chỉ điền cho kiến trúc có artifact train/deploy hiện có trong project.

![Architecture ablation]({arch_fig.as_posix()})

{chr(10).join(arch_lines)}

## Train đồng nhất baseline kiến trúc

Các baseline đã được train trong cùng một lệnh với cùng split, seed, optimizer, 10 epochs, image size 96 và loss profile dark-map. Đây là run ablation ngắn đồng nhất để so công bằng ban đầu; không thay thế model refine dài hạn.

![Unified architecture training]({arch_train_fig.as_posix()})

{chr(10).join(arch_train_lines)}

## Thông số mô hình triển khai trên STM32

| Hạng mục | Giá trị |
|---|---:|
| Input | `input_rgb`, uint8, 1x3x96x96 |
| Output | `enhanced_rgb`, uint8, 1x3x96x96 |
| MACC | 19,386,115 |
| Weights | 5.79 KiB |
| Activations | 518.96 KiB |
| Layout khớp PC-board | raw NCHW |
| PC-board exact match | 98.09% |
| PC-board MAE lượng tử | 0.022 / 255 |
| PC-board cosine | 1.000000 |

## Chi phí inference trên board

| Metric | Trung bình |
|---|---:|
| Camera/display | {fmt(board['camera_display_ms']['mean'])} ms/frame |
| Preprocess | {fmt(board['preprocess_ms']['mean'])} ms/frame |
| Inference | {fmt(board['inference_ms']['mean'])} ms/frame |
| Postprocess | {fmt(board['postprocess_ms']['mean'])} ms/frame |
| Total pipeline | {fmt(board['total_ms']['mean'])} ms/frame |
| FPS | {fmt(board['fps']['mean'])} |

## So sánh thị giác qua board

![Board visual]({contact.as_posix()})

{chr(10).join(visual_lines)}

![Histogram]({hist.as_posix()})

## Nhận xét kỹ thuật

- PC ONNX và board output khớp gần như tuyệt đối, nên lỗi sai khác hiện tại không còn nằm ở layout PC-board với dump đã kiểm chứng.
- Chất lượng ảnh cuối phụ thuộc chủ yếu vào student ONNX đang deploy và postprocess hiển thị.
- Framework training có teacher/loss/Optuna, nhưng artifact firmware là student nhỏ. Vì vậy báo cáo/paper nên ghi rõ hai pha: training có các thành phần nâng cao; deployment chỉ giữ student tối ưu hóa.
- Nếu cần số paper cuối cùng, nên chạy một lần retrain/export/benchmark đồng nhất từ cùng seed, dataset split và checkpoint, rồi cập nhật lại bảng này.
"""
    OUT_MD.write_text(text, encoding="utf-8")


def write_docx(data: dict) -> None:
    best = data["best"]
    board = data["board_timing"]
    pc = data["pc"]
    visual = data["visual"]
    arch = data["arch"]
    arch_train = data["arch_train"]

    doc = Document()
    apply_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DarkGhost-ESPNet")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dark-map Guided Ghost-ESP Network for Real-time Low-light Image Enhancement on STM32")
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string("404040")

    doc.add_paragraph("Báo cáo benchmark đồng bộ theo log/train/PC ONNX/board hiện có trong project, ngày 2026-07-26.")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Báo cáo này tổng hợp mô hình DarkGhost-ESPNet, một student nhẹ dựa trên Ghost-ESP cho bài toán "
        "low-light image enhancement 96x96. Framework huấn luyện kết hợp dark-map guidance, teacher RGB "
        "Retinex/VD, adaptive loss theo vùng tối và Optuna để tinh chỉnh trọng số loss. Trên STM32, firmware "
        "chỉ triển khai student DarkGhost-ESPNet-Tiny96 qua Cube.AI INT8/u8 runtime."
    )

    doc.add_heading("1. Proposed Method", level=1)
    add_figure(
        doc,
        FIG_DARK / "darkghost_espnet_training_deployment_framework.png",
        "Figure 1. DarkGhost-ESPNet training and deployment framework.",
        6.7,
    )
    doc.add_paragraph(
        "Trong pha training, ảnh low-light được đưa qua student DarkGhost-ESPNet và đồng thời tạo dark-map "
        "để nhấn mạnh vùng thiếu sáng. Teacher Retinex/VD cung cấp tín hiệu tham chiếu RGB. Loss tổng hợp gồm "
        "Charbonnier, SSIM, perceptual, color/chroma và thành phần điều biến theo dark-map. Optuna được dùng "
        "để tìm cấu hình trọng số loss/siêu tham số. Trong pha deployment, chỉ student đã export ONNX/Cube.AI "
        "được giữ lại; teacher, loss và Optuna không chạy trên board."
    )
    add_figure(
        doc,
        FIG_MODEL / "serious_current_method_ghost_esp_darkmap_retinex_optuna.png",
        "Figure 2. Tóm tắt các thành phần phù hợp với project: Ghost-ESP, dark-map loss, teacher Retinex/VD và Optuna.",
        6.7,
    )

    doc.add_heading("2. Số liệu train mới nhất", level=1)
    add_table(
        doc,
        ["Hạng mục", "Giá trị"],
        [
            ["Experiment", best["experiment"]],
            ["Epochs", best["epochs"]],
            ["Best PSNR", f"{fmt(best['best_psnr'])} dB tại epoch {best['best_psnr_epoch']}"],
            ["Best SSIM", f"{fmt(best['best_ssim'])} tại epoch {best['best_ssim_epoch']}"],
            ["Epoch cuối", best["last_epoch"]],
            ["Val PSNR epoch cuối", fmt(best["last_val_psnr"])],
            ["Val SSIM epoch cuối", fmt(best["last_val_ssim"])],
            ["Teacher/reference PSNR", fmt(best["teacher_psnr"])],
            ["Teacher/reference SSIM", fmt(best["teacher_ssim"])],
        ],
    )
    doc.add_paragraph(
        "Các số trên lấy từ train log mới nhất đã đồng bộ trong thư mục experiments/refine. Vì vậy báo cáo "
        "không lấy epoch 80 cũ làm kết quả chính."
    )

    doc.add_heading("3. PC ONNX benchmark theo dataset split", level=1)
    rows = []
    for name, row in pc.items():
        rows.append(
            [
                name,
                str(row["n"]),
                fmt(row["psnr_mean"]),
                fmt(row["ssim_mean"]),
                fmt(row["mae_mean"], 6),
                fmt(row["inference_ms_pc_mean"]),
            ]
        )
    add_table(doc, ["Dataset split", "N", "PSNR", "SSIM", "MAE", "PC ONNX ms/frame"], rows)

    doc.add_heading("4. Benchmark kiến trúc model", level=1)
    doc.add_paragraph(
        "Bảng này so sánh ablation kiến trúc tương tự paper trên input 96x96. Các chỉ số chi phí được đo trực tiếp "
        "bằng PyTorch CPU. PSNR/SSIM/MAE chỉ được ghi khi project có checkpoint/ONNX đã train tương ứng."
    )
    add_figure(
        doc,
        FIG_ARCH / "architecture_cost_comparison.png",
        "Figure 3. So sánh chi phí kiến trúc Conv2D, Separable, GhostSep, Ghost-ESP và DarkGhost-ESPNet.",
        6.7,
    )
    arch_table_rows = []
    for row in arch:
        arch_table_rows.append(
            [
                row["architecture"],
                f"{int(row['params']):,}",
                row["model_size_fp32_kib"],
                row["model_size_int8_est_kib"],
                f"{int(row['macs']):,}",
                row["latency_ms_mean"],
                row["psnr"] or "-",
                row["ssim"] or "-",
                row["quality_source"],
            ]
        )
    add_table(
        doc,
        ["Architecture", "Params", "FP32 KiB", "INT8 est. KiB", "MACs", "CPU ms", "PSNR", "SSIM", "Source"],
        arch_table_rows,
    )
    doc.add_paragraph(
        "Các baseline Conv2D/Separable/GhostSep/Ghost-ESP hiện mới được benchmark chi phí vì project chưa còn "
        "checkpoint train đồng nhất cho chúng. Để có PSNR/SSIM công bằng như paper, cần retrain từng kiến trúc với "
        "cùng seed, split, epoch và loss."
    )

    doc.add_heading("5. Train đồng nhất baseline kiến trúc", level=1)
    doc.add_paragraph(
        "Các baseline được train lại trong cùng một lệnh, cùng split LOL + LOL-v2-Real, cùng seed 42, image size 96, "
        "optimizer AdamW, 10 epochs và loss profile dark-map. Đây là run ablation ngắn để có số liệu đồng nhất ban đầu; "
        "model refine dài hạn hiện tại vẫn là kết quả tốt hơn cho deployment."
    )
    add_figure(
        doc,
        FIG_ARCH_TRAIN / "unified_architecture_quality_comparison.png",
        "Figure 4. Chất lượng PSNR/SSIM/MAE sau run train đồng nhất 10 epochs.",
        6.7,
    )
    arch_train_table_rows = []
    for row in arch_train:
        arch_train_table_rows.append(
            [
                row["architecture"],
                f"{int(row['params']):,}",
                f"{int(row['macs']):,}",
                row["best_epoch"],
                fmt(row["best_psnr"]),
                fmt(row["best_ssim"]),
                fmt(row["best_mae"], 6),
            ]
        )
    add_table(
        doc,
        ["Architecture", "Params", "MACs", "Best epoch", "PSNR", "SSIM", "MAE"],
        arch_train_table_rows,
    )
    doc.add_paragraph(
        "Trong run ngắn này, GhostSep đang cho PSNR/SSIM tốt nhất. DarkGhost-ESPNet chưa vượt ở 10 epochs, "
        "vì lợi thế dark-map/distillation thường cần lịch train/refine dài hơn để hội tụ."
    )

    doc.add_heading("6. Model triển khai và tương đương PC-board", level=1)
    add_table(
        doc,
        ["Hạng mục", "Giá trị"],
        [
            ["ONNX deploy", "ghost_esp_dark_w12_m24_enhanced_rgb_u8out_tail_simplified_nchw.onnx"],
            ["Input", "input_rgb, uint8, 1x3x96x96"],
            ["Output", "enhanced_rgb, uint8, 1x3x96x96"],
            ["MACC", "19,386,115"],
            ["Weights", "5.79 KiB"],
            ["Activations", "518.96 KiB"],
            ["Best layout PC-board", "raw NCHW"],
            ["Exact match", "98.09%"],
            ["MAE lượng tử", "0.022 / 255"],
            ["Cosine", "1.000000"],
        ],
    )

    doc.add_heading("7. Chi phí inference trên board", level=1)
    add_table(
        doc,
        ["Metric", "Trung bình"],
        [
            ["Camera/display", f"{fmt(board['camera_display_ms']['mean'])} ms/frame"],
            ["Preprocess", f"{fmt(board['preprocess_ms']['mean'])} ms/frame"],
            ["Inference", f"{fmt(board['inference_ms']['mean'])} ms/frame"],
            ["Postprocess", f"{fmt(board['postprocess_ms']['mean'])} ms/frame"],
            ["Total pipeline", f"{fmt(board['total_ms']['mean'])} ms/frame"],
            ["FPS", fmt(board["fps"]["mean"])],
        ],
    )

    doc.add_heading("8. So sánh thị giác qua board", level=1)
    add_figure(doc, FIG_REAL / "board_input_vs_ai_output_contact_x4.png", "Figure 5. Input/preprocess và board AI output.", 6.5)
    rows = []
    for key in [
        "brightness_mean",
        "contrast_std",
        "saturation_mean",
        "sharpness_laplacian_abs_mean",
        "clip_0_ratio_rgb",
        "clip_255_ratio_rgb",
    ]:
        rows.append([key, fmt(visual["input_preprocess"][key], 6), fmt(visual["ai_output"][key], 6)])
    add_table(doc, ["Metric", "Input/preprocess", "Board AI output"], rows)
    add_figure(doc, FIG_REAL / "luma_histogram_input_vs_ai_output.png", "Figure 6. Histogram luminance input so với AI output.", 6.2)

    doc.add_heading("9. Discussion", level=1)
    add_bullets(
        doc,
        [
            "PC ONNX và board output khớp rất sát theo dump hiện tại, nên sai khác PC-board không còn là nguyên nhân chính.",
            "Teacher Retinex/VD, adaptive loss và Optuna là đóng góp ở pha training; firmware chỉ chạy student nhỏ.",
            "Để có số liệu paper cuối cùng, nên chạy lại một pipeline retrain-export-benchmark thống nhất từ cùng seed, split và checkpoint.",
        ],
    )

    doc.add_heading("10. Conclusion", level=1)
    doc.add_paragraph(
        "DarkGhost-ESPNet là tên phù hợp cho hướng đóng góp hiện tại: Ghost-ESP nhẹ, có guidance bằng dark-map, "
        "loss thích nghi theo vùng tối và teacher Retinex/VD trong training. Biến thể deploy DarkGhost-ESPNet-Tiny96 "
        "đạt khoảng 170.3 ms inference/frame và khoảng 191.6 ms toàn pipeline trên board theo log hiện có."
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    BENCH_DARK.mkdir(parents=True, exist_ok=True)
    data = collect_data()
    write_markdown(data)
    write_docx(data)
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
