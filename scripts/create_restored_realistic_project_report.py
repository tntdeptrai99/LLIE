from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATE = "20260726"
OUT_DIR = ROOT / "reports" / "benchmarks" / f"restored_realistic_project_{DATE}"
OUT_MD = OUT_DIR / "bao_cao_so_lieu_thuc_te_sau_khoi_phuc.md"
OUT_DOCX = OUT_DIR / "bao_cao_so_lieu_thuc_te_sau_khoi_phuc_v3_paper_figures.docx"

METRICS_DIR = ROOT / "reports" / "metrics"
CURRENT_MODEL = ROOT / "reports" / "benchmarks" / f"current_model_{DATE}"
CURRENT_PIPE = ROOT / "reports" / "benchmarks" / f"current_pipeline_{DATE}"
REAL_FIG = ROOT / "reports" / "figures" / f"current_pipeline_{DATE}" / "real_camera_visual"
METHOD_FIG = ROOT / "reports" / "figures" / f"darkghost_espnet_{DATE}" / "darkghost_espnet_training_deployment_framework.png"
ARCH_SAT = ROOT / "reports" / "benchmarks" / f"architecture_ablation_saturation_{DATE}" / "best_available_saturation_summary.csv"
COMPONENT_ABLATION = ROOT / "reports" / "benchmarks" / f"component_ablation_{DATE}" / "component_ablation_summary.csv"
COMPONENT_FIG = ROOT / "reports" / "figures" / f"component_ablation_{DATE}" / "component_ablation_effect.png"
PAPER_FIG = ROOT / "reports" / "figures" / f"paper_style_{DATE}"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def parse_summary_txt(path: Path) -> dict[str, str]:
    data: dict[str, str] = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        if "=" in line:
            key, value = line.split("=", 1)
            data[key.strip()] = value.strip()
    return data


def load_best_research_summaries() -> list[dict[str, str]]:
    selected_names = [
        "ghost_esp_dark_w12_m24_gain3_res035_96_best",
        "ghost_esp_dark_w12_m24_gain3_res035_optuna_trial011_best",
        "ghost_esp_dark_w12_m24_gain3_res035_trial011_long80_best_ssim",
        "ghost_esp_dark_w12_m24_gain3_res035_plateau_score_best_monitor",
        "ghost_esp_dark_w12_m24_gain3_res035_plateau_score_best_ssim",
    ]
    rows = []
    for target in selected_names:
        matches = sorted(METRICS_DIR.glob(f"lol_test_{target}_{target}_summary.txt"))
        if not matches:
            matches = sorted(METRICS_DIR.glob(f"*{target}*summary.txt"))
        if not matches:
            continue
        data = parse_summary_txt(matches[0])
        rows.append(
            {
                "model": target,
                "count": data.get("count", ""),
                "psnr": data.get("psnr", ""),
                "ssim": data.get("ssim", ""),
                "source": str(matches[0]),
            }
        )
    return rows


def fmt(value: str | float, digits: int = 4) -> str:
    return f"{float(value):.{digits}f}"


def shade(cell, color: str = "EAF2F8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C8D6E5")


def set_cell(cell, value, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(8.8)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    borders(table)
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell)
        set_cell(cell, header, True, True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell(cell, value, center=len(str(value)) < 18)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.5) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("555555")


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in [
        ("Heading 1", 15, "17365D"),
        ("Heading 2", 12, "1F4E79"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def collect() -> dict:
    research = load_best_research_summaries()
    pc = json.loads((CURRENT_MODEL / "pc_dataset_split_summary.json").read_text(encoding="utf-8"))
    timing = {row["metric"]: row for row in read_csv(CURRENT_MODEL / "board_timing_summary.csv")}
    visual = {row["name"]: row for row in read_csv(REAL_FIG / "image_metrics.csv")}
    arch = read_csv(ARCH_SAT) if ARCH_SAT.exists() else []
    component = read_csv(COMPONENT_ABLATION) if COMPONENT_ABLATION.exists() else []
    return {
        "research": research,
        "pc": pc,
        "timing": timing,
        "visual": visual,
        "arch": arch,
        "component": component,
    }


def write_markdown(data: dict) -> None:
    research_lines = [
        "| Model thực tế trong project | N | PSNR | SSIM |",
        "|---|---:|---:|---:|",
    ]
    for row in data["research"]:
        research_lines.append(f"| `{row['model']}` | {row['count']} | {fmt(row['psnr'])} | {fmt(row['ssim'])} |")

    pc_lines = [
        "| Dataset split | N | PSNR | SSIM | MAE | PC ONNX ms/frame |",
        "|---|---:|---:|---:|---:|---:|",
    ]
    for name, row in data["pc"].items():
        pc_lines.append(
            f"| {name} | {row['n']} | {fmt(row['psnr_mean'])} | {fmt(row['ssim_mean'])} | "
            f"{fmt(row['mae_mean'], 6)} | {fmt(row['inference_ms_pc_mean'])} |"
        )

    arch_lines = [
        "| Architecture | Status | Epochs ran | Best epoch | PSNR | SSIM | MAE |",
        "|---|---|---:|---:|---:|---:|---:|",
    ]
    for row in data["arch"]:
        arch_lines.append(
            f"| {row['architecture']} | {row['status']} | {row['epochs_ran']} | {row['best_epoch']} | "
            f"{fmt(row['best_psnr'])} | {fmt(row['best_ssim'])} | {fmt(row['best_mae'], 6)} |"
        )

    component_lines = [
        "| Variant | N | Epoch/source | PSNR | SSIM | Notes |",
        "|---|---:|---|---:|---:|---|",
    ]
    for row in data["component"]:
        component_lines.append(
            f"| {row['variant']} | {row['count']} | {row['epoch']} | "
            f"{fmt(row['psnr'])} | {fmt(row['ssim'])} | {row['notes']} |"
        )

    timing = data["timing"]
    visual = data["visual"]
    text = f"""# Báo cáo số liệu thực tế sau khi khôi phục project

Ngày tạo: 2026-07-26

## Phạm vi

Báo cáo này chỉ lấy các số liệu có artifact thật trong project sau khi khôi phục: checkpoint/log Ghost-ESP/DarkGhost, ONNX đang deploy, PC-board dump, board timing, visual board metrics và ablation architecture đã chạy.

## Best research checkpoint đã khôi phục

{chr(10).join(research_lines)}

Kết luận: checkpoint nghiên cứu tốt nhất trong artifact khôi phục là `plateau_score_best_monitor`, đạt PSNR 19.6221 và SSIM 0.843353 trên `splits/lol_test.txt` với 15 ảnh.

## Current deployed ONNX

ONNX đang dùng cho firmware hiện tại: `stm32/onnx/ghost_esp_dark_w12_m24_enhanced_rgb_u8out_tail_simplified_nchw.onnx`.

{chr(10).join(pc_lines)}

## Board runtime hiện tại

| Metric | Mean |
|---|---:|
| Camera/LCD display | {fmt(timing['camera_display_ms']['mean'])} ms |
| Preprocess | {fmt(timing['preprocess_ms']['mean'])} ms |
| Inference | {fmt(timing['inference_ms']['mean'])} ms |
| Postprocess | {fmt(timing['postprocess_ms']['mean'])} ms |
| Total pipeline | {fmt(timing['total_ms']['mean'])} ms |
| FPS | {fmt(timing['fps']['mean'])} |

PC-board equivalence đã kiểm chứng: layout `raw_nchw`, exact match 98.09%, MAE lượng tử 0.022/255, cosine 1.000000.

## Visual board metrics

![Board input/output]({(REAL_FIG / 'board_input_vs_ai_output_contact_x4.png').as_posix()})

| Metric | Input/preprocess | Board AI output |
|---|---:|---:|
| Brightness mean | {fmt(visual['input_preprocess']['brightness_mean'], 6)} | {fmt(visual['ai_output']['brightness_mean'], 6)} |
| Contrast std | {fmt(visual['input_preprocess']['contrast_std'], 6)} | {fmt(visual['ai_output']['contrast_std'], 6)} |
| Saturation mean | {fmt(visual['input_preprocess']['saturation_mean'], 6)} | {fmt(visual['ai_output']['saturation_mean'], 6)} |
| Sharpness Laplacian abs mean | {fmt(visual['input_preprocess']['sharpness_laplacian_abs_mean'], 6)} | {fmt(visual['ai_output']['sharpness_laplacian_abs_mean'], 6)} |
| Clip 0 RGB | {fmt(visual['input_preprocess']['clip_0_ratio_rgb'], 6)} | {fmt(visual['ai_output']['clip_0_ratio_rgb'], 6)} |
| Clip 255 RGB | {fmt(visual['input_preprocess']['clip_255_ratio_rgb'], 6)} | {fmt(visual['ai_output']['clip_255_ratio_rgb'], 6)} |

## Architecture ablation thực tế

{chr(10).join(arch_lines)}

Lưu ý: run saturation bị ngắt, nên Ghost-ESP chỉ có partial 4 epoch và DarkGhost-ESPNet chưa có trong run này. Bảng này dùng để tham khảo trạng thái đã ghi được, không thay thế long-refine model.

## Component ablation: Optuna, dark-map loss, teacher

![Component ablation]({COMPONENT_FIG.as_posix()})

{chr(10).join(component_lines)}

Ghi chú: `không có loss` ở đây được hiểu là không dùng proposed adaptive dark-map loss; vẫn phải có reconstruction/GT loss để model học được. Hai run `No adaptive dark-map loss` và `No teacher KD` là benchmark CUDA 30 epoch mới, còn các dòng Optuna/proposed là artifact long-run đã khôi phục, vì vậy cần ghi rõ ngân sách train khi dùng trong paper.
"""
    OUT_MD.write_text(text, encoding="utf-8")


def write_docx(data: dict) -> None:
    doc = Document()
    apply_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Báo cáo số liệu thực tế sau khi khôi phục project")
    r.bold = True
    r.font.size = Pt(21)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Ghost-ESP / DarkGhost-ESPNet | artifact thực tế đang có | 2026-07-26")
    r.italic = True
    r.font.size = Pt(11)

    doc.add_heading("1. Phạm vi và nguyên tắc chọn số liệu", level=1)
    doc.add_paragraph(
        "Báo cáo chỉ dùng số liệu có artifact thật sau khi khôi phục: train_log/checkpoint Ghost-ESP/DarkGhost, "
        "metric summary đã lưu, ONNX đang deploy, PC-board dump, board timing và visual board metrics. Các baseline "
        "không liên quan trực tiếp hoặc không có checkpoint/ONNX thật không được dùng làm kết luận chính."
    )
    add_figure(doc, METHOD_FIG, "Figure 1. Framework DarkGhost-ESPNet dùng trong project.", 6.6)

    doc.add_heading("2. Best research checkpoint đã khôi phục", level=1)
    add_table(
        doc,
        ["Model thực tế trong project", "N", "PSNR", "SSIM"],
        [[row["model"], row["count"], fmt(row["psnr"]), fmt(row["ssim"])] for row in data["research"]],
    )
    doc.add_paragraph(
        "Checkpoint tốt nhất trong artifact khôi phục là plateau_score_best_monitor, đạt PSNR 19.6221 và SSIM "
        "0.843353 trên splits/lol_test.txt với 15 ảnh. Đây là best research checkpoint, không đồng nghĩa tự động "
        "là đúng file đang flash trên board hiện tại."
    )

    doc.add_heading("3. Current deployed ONNX", level=1)
    doc.add_paragraph(
        "ONNX đang dùng cho firmware hiện tại: "
        "stm32/onnx/ghost_esp_dark_w12_m24_enhanced_rgb_u8out_tail_simplified_nchw.onnx."
    )
    pc_rows = []
    for name, row in data["pc"].items():
        pc_rows.append(
            [
                name,
                row["n"],
                fmt(row["psnr_mean"]),
                fmt(row["ssim_mean"]),
                fmt(row["mae_mean"], 6),
                fmt(row["inference_ms_pc_mean"]),
            ]
        )
    add_table(doc, ["Dataset split", "N", "PSNR", "SSIM", "MAE", "PC ONNX ms/frame"], pc_rows)

    doc.add_heading("4. Board runtime và PC-board equivalence", level=1)
    timing = data["timing"]
    add_table(
        doc,
        ["Metric", "Mean"],
        [
            ["Camera/LCD display", f"{fmt(timing['camera_display_ms']['mean'])} ms"],
            ["Preprocess", f"{fmt(timing['preprocess_ms']['mean'])} ms"],
            ["Inference", f"{fmt(timing['inference_ms']['mean'])} ms"],
            ["Postprocess", f"{fmt(timing['postprocess_ms']['mean'])} ms"],
            ["Total pipeline", f"{fmt(timing['total_ms']['mean'])} ms"],
            ["FPS", fmt(timing["fps"]["mean"])],
        ],
    )
    doc.add_paragraph(
        "PC-board equivalence đã kiểm chứng trên dump hiện tại: layout raw_nchw, exact match 98.09%, "
        "MAE lượng tử 0.022/255, max abs 3 và cosine 1.000000."
    )

    doc.add_heading("5. Visual board metrics", level=1)
    add_figure(doc, REAL_FIG / "board_input_vs_ai_output_contact_x4.png", "Figure 2. Board input/preprocess và AI output.", 6.4)
    visual = data["visual"]
    add_table(
        doc,
        ["Metric", "Input/preprocess", "Board AI output"],
        [
            ["Brightness mean", fmt(visual["input_preprocess"]["brightness_mean"], 6), fmt(visual["ai_output"]["brightness_mean"], 6)],
            ["Contrast std", fmt(visual["input_preprocess"]["contrast_std"], 6), fmt(visual["ai_output"]["contrast_std"], 6)],
            ["Saturation mean", fmt(visual["input_preprocess"]["saturation_mean"], 6), fmt(visual["ai_output"]["saturation_mean"], 6)],
            ["Sharpness Laplacian abs mean", fmt(visual["input_preprocess"]["sharpness_laplacian_abs_mean"], 6), fmt(visual["ai_output"]["sharpness_laplacian_abs_mean"], 6)],
            ["Clip 0 RGB", fmt(visual["input_preprocess"]["clip_0_ratio_rgb"], 6), fmt(visual["ai_output"]["clip_0_ratio_rgb"], 6)],
            ["Clip 255 RGB", fmt(visual["input_preprocess"]["clip_255_ratio_rgb"], 6), fmt(visual["ai_output"]["clip_255_ratio_rgb"], 6)],
        ],
    )
    add_figure(doc, REAL_FIG / "luma_histogram_input_vs_ai_output.png", "Figure 3. Histogram luminance input và AI output.", 6.1)

    doc.add_heading("6. Architecture ablation thực tế", level=1)
    add_table(
        doc,
        ["Architecture", "Status", "Epochs", "Best epoch", "PSNR", "SSIM", "MAE"],
        [
            [
                row["architecture"],
                row["status"],
                row["epochs_ran"],
                row["best_epoch"],
                fmt(row["best_psnr"]),
                fmt(row["best_ssim"]),
                fmt(row["best_mae"], 6),
            ]
            for row in data["arch"]
        ],
    )
    doc.add_paragraph(
        "Run saturation bị ngắt, nên Ghost-ESP chỉ có partial 4 epoch và DarkGhost-ESPNet chưa có trong run này. "
        "Bảng này là trạng thái đã ghi được, không thay thế kết quả long-refine của checkpoint chính."
    )

    doc.add_heading("7. Component ablation: Optuna, dark-map loss, teacher", level=1)
    add_figure(doc, COMPONENT_FIG, "Figure 4. Hiệu quả các thành phần Optuna, adaptive dark-map loss và teacher KD.", 6.5)
    add_table(
        doc,
        ["Variant", "N", "Epoch/source", "PSNR", "SSIM", "Notes"],
        [
            [
                row["variant"],
                row["count"],
                row["epoch"],
                fmt(row["psnr"]),
                fmt(row["ssim"]),
                row["notes"],
            ]
            for row in data["component"]
        ],
    )
    doc.add_paragraph(
        "Trong báo cáo này, 'không có loss' được hiểu là không dùng proposed adaptive dark-map loss; vẫn cần "
        "reconstruction/GT loss để model học được. Hai run No adaptive dark-map loss và No teacher KD là benchmark "
        "CUDA 30 epoch mới. Các dòng Optuna/proposed là artifact long-run đã khôi phục, nên khi đưa vào paper cần ghi rõ "
        "ngân sách train khác nhau."
    )

    doc.add_heading("8. Paper-style figures", level=1)
    add_figure(
        doc,
        PAPER_FIG / "optuna_optimization_progress.png",
        "Figure 5. Optuna optimization progress theo objective PSNR + 5 x SSIM trên 20 trials thật trong project.",
        6.5,
    )
    add_figure(
        doc,
        PAPER_FIG / "optuna_best_loss_weights.png",
        "Figure 6. Bộ trọng số loss tốt nhất do Optuna chọn ở trial 011.",
        6.3,
    )
    add_figure(
        doc,
        PAPER_FIG / "training_validation_curves_project.png",
        "Figure 7. Training/validation curves từ train_log. Các curve này dùng validation trong quá trình train, khác với bảng eval summary trên lol_test.",
        6.7,
    )
    add_figure(
        doc,
        PAPER_FIG / "component_ablation_training_curves.png",
        "Figure 8. Curves 30 epoch cho ablation có kiểm soát: tắt adaptive dark-map loss và tắt teacher KD.",
        6.7,
    )

    doc.add_heading("9. Kết luận dùng cho báo cáo/paper", level=1)
    doc.add_paragraph(
        "Số nên dùng làm kết quả nghiên cứu chính: DG-GhostESP-96 plateau best_monitor với PSNR 19.6221 và SSIM "
        "0.843353 trên lol_test. Số nên dùng làm kết quả deployment hiện tại: ONNX tail_simplified_nchw, PC-board "
        "match rất sát và board chạy khoảng 170.3 ms inference/frame, tổng pipeline khoảng 191.6 ms/frame."
    )
    doc.save(OUT_DOCX)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = collect()
    write_markdown(data)
    write_docx(data)
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
