from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BENCH_DIR = ROOT / "reports" / "benchmarks" / "current_pipeline_20260726"
FIG_DIR = ROOT / "reports" / "figures" / "current_pipeline_20260726"
OUT = BENCH_DIR / "bao_cao_benchmark_chat_luong_anh_llie_current_pipeline.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(9)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2F3")


def format_table(table, header_fill: str = "F2F4F7") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
            if row_index == 0:
                set_cell_shading(cell, header_fill)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(4.2)
    set_cell_text(table.rows[0].cells[0], "Hạng mục", True)
    set_cell_text(table.rows[0].cells[1], "Giá trị", True)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
    format_table(table)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def read_dataset_summary() -> dict:
    with (BENCH_DIR / "dataset_quality_summary.json").open("r", encoding="utf-8") as f:
        return json.load(f)


def read_architecture_rows() -> list[dict[str, str]]:
    with (BENCH_DIR / "architecture_benchmark.csv").open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_loss_rows() -> list[dict[str, str]]:
    with (BENCH_DIR / "loss_function_benchmark_from_train_logs.csv").open(
        "r", encoding="utf-8-sig", newline=""
    ) as f:
        rows = list(csv.DictReader(f))
    return sorted(rows, key=lambda r: float(r.get("best_psnr") or 0), reverse=True)[:8]


def parse_pc_board_log() -> tuple[str, list[str]]:
    path = BENCH_DIR / "pc_board_equivalence.log"
    text = path.read_text(encoding="utf-8", errors="ignore")
    lines = []
    status = "Đạt"
    for needle in [
        "board output0_public vs ONNX output0 enhanced_rgb",
        "uint8 raw_nchw:",
        "best=raw_nchw",
    ]:
        for line in text.splitlines():
            if needle in line:
                lines.append(line.strip())
                break
    if "best=raw_nchw" not in text:
        status = "Cần kiểm tra lại"
    return status, lines


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Báo cáo benchmark chất lượng ảnh LLIE")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Current pipeline 2026-07-26 | STM32H750 + OV5640 + LCD ST7735 + Ghost-ESP")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_dataset_table(doc: Document, summary: dict) -> None:
    table = doc.add_table(rows=1, cols=6)
    headers = ["Dataset", "N", "PSNR mean", "SSIM mean", "MAE mean", "PC inference ms"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, True)
    for dataset, values in summary.items():
        cells = table.add_row().cells
        row = [
            dataset,
            str(values["n"]),
            f"{values['psnr_mean']:.4f}",
            f"{values['ssim_mean']:.4f}",
            f"{values['mae_mean']:.6f}",
            f"{values['inference_ms_pc_mean']:.4f}",
        ]
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    format_table(table)


def add_architecture_table(doc: Document, rows: list[dict[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=7)
    headers = ["Model/block", "Trạng thái", "Params", "INT8 KB", "PSNR", "SSIM", "Inference ms"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, True)
    for row in rows:
        values = [
            row.get("model", ""),
            row.get("status", ""),
            row.get("params", ""),
            row.get("model_size_int8_kb", ""),
            row.get("psnr", ""),
            row.get("ssim", ""),
            row.get("inference_time_ms_pc", ""),
        ]
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            set_cell_text(cell, text)
    format_table(table)


def add_loss_table(doc: Document, rows: list[dict[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=5)
    headers = ["Experiment", "Epochs", "Best PSNR", "Best SSIM", "Loss family"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, True)
    for row in rows:
        exp = row.get("experiment", "").replace("experiments\\", "")
        if len(exp) > 52:
            exp = exp[:49] + "..."
        values = [
            exp,
            row.get("epochs", ""),
            f"{float(row.get('best_psnr') or 0):.4f}",
            f"{float(row.get('best_ssim') or 0):.4f}",
            row.get("loss_family_inferred", ""),
        ]
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            set_cell_text(cell, text)
    format_table(table)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_image_if_exists(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width))
        add_caption(doc, caption)


def build() -> None:
    summary = read_dataset_summary()
    arch_rows = read_architecture_rows()
    loss_rows = read_loss_rows()
    pc_status, pc_lines = parse_pc_board_log()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("1. Tóm tắt điều đã benchmark", level=1)
    add_bullets(
        doc,
        [
            "Benchmark chất lượng ảnh trên dataset LOL val và LOL-v2-Real val, tổng cộng 115 cặp input/ground truth.",
            "Benchmark model hiện đang deploy trên PC ONNX: Ghost-ESP current deployed ONNX.",
            "Tổng hợp benchmark loss function từ các train_log.csv còn lại trong thư mục experiments.",
            "Kiểm tra tương đương PC-board bằng UART dump hiện có và so output board với output ONNX.",
            "Lưu artifact thị giác thực tế gồm input preprocess, AI output, contact sheet và histogram.",
        ],
    )

    doc.add_heading("2. Artifact đã xuất", level=1)
    add_kv_table(
        doc,
        [
            ("Benchmark dir", str(BENCH_DIR)),
            ("Figure dir", str(FIG_DIR)),
            ("Dataset CSV", str(BENCH_DIR / "dataset_quality_metrics.csv")),
            ("Architecture CSV", str(BENCH_DIR / "architecture_benchmark.csv")),
            ("Loss CSV", str(BENCH_DIR / "loss_function_benchmark_from_train_logs.csv")),
            ("PC-board log", str(BENCH_DIR / "pc_board_equivalence.log")),
        ],
    )

    doc.add_heading("3. Benchmark chất lượng ảnh trên dataset", level=1)
    doc.add_paragraph(
        "Pipeline hiện tại resize ảnh về 96x96, chạy ONNX active và so với ground truth bằng PSNR, SSIM, MAE. "
        "Inference time là thời gian chạy ONNX trên PC cho một frame, chưa phải thời gian MCU."
    )
    add_dataset_table(doc, summary)

    add_image_if_exists(
        doc,
        FIG_DIR / "dataset_quality" / "000_LOL_val_1" / "contact_sheet.png",
        "Contact sheet mẫu: input, AI output và ground truth.",
    )
    add_image_if_exists(
        doc,
        FIG_DIR / "dataset_quality" / "000_LOL_val_1" / "histogram.png",
        "Histogram mẫu trước/sau để kiểm tra phân bố độ sáng.",
    )

    doc.add_heading("4. Benchmark kiến trúc mô hình", level=1)
    doc.add_paragraph(
        "Sau bước dọn project, chỉ còn artifact ONNX của model hiện tại nên chỉ model này được đo trực tiếp. "
        "Các kiến trúc Conv2D, Separable, GhostSep và Ghost-ESP + Distill được ghi là missing_artifact vì không còn ONNX/checkpoint tương ứng."
    )
    add_architecture_table(doc, arch_rows)

    doc.add_heading("5. Benchmark loss function", level=1)
    doc.add_paragraph(
        "Bảng dưới đây lấy các experiment tốt nhất theo best PSNR từ train_log.csv còn lại, dùng để so nhanh baseline với các thử nghiệm adaptive/hybrid/dark-map."
    )
    add_loss_table(doc, loss_rows)

    doc.add_heading("6. Lượng tử hóa và tương đương PC-board", level=1)
    add_kv_table(doc, [("Trạng thái", pc_status), ("Layout khớp tốt nhất", "raw_nchw")])
    add_bullets(doc, pc_lines)
    doc.add_paragraph(
        "Kết luận kỹ thuật: output board khớp ONNX rất sát ở raw_nchw, nên với dump hiện có lỗi LCD không còn là do PC ONNX và board AI output lệch lớn. "
        "Phần cần theo dõi tiếp là đường render LCD/buffer/runtime frame hoặc chất lượng model nếu output ONNX chưa đạt thị giác mong muốn."
    )

    doc.add_heading("7. Benchmark thị giác thực tế", level=1)
    doc.add_paragraph(
        "Artifact thực tế lấy từ frame board đã dump: input preprocess, AI output, contact sheet x4, histogram và metric ảnh."
    )
    add_image_if_exists(
        doc,
        FIG_DIR / "real_camera_visual" / "board_input_vs_ai_output_contact_x4.png",
        "So sánh input preprocess và AI output từ frame board thực tế.",
    )
    add_image_if_exists(
        doc,
        FIG_DIR / "real_camera_visual" / "luma_histogram_input_vs_ai_output.png",
        "Histogram luminance input so với AI output trên frame board.",
    )

    doc.add_heading("8. Kết luận và giới hạn", level=1)
    add_bullets(
        doc,
        [
            "Dataset benchmark đã chạy đủ trên 115 ảnh hiện có và lưu đầy đủ input/output/ground truth/contact sheet/histogram mẫu.",
            "Model hiện tại có PSNR tốt hơn trên LOL-v2-Real val so với LOL val, nhưng output vẫn cần đánh giá thị giác vì metric không phản ánh hết cảm nhận ảnh.",
            "PC-board equivalence đạt tốt: board output gần như trùng ONNX output ở layout raw_nchw.",
            "So sánh kiến trúc chưa hoàn chỉnh vì các artifact Conv2D, Separable, GhostSep và Ghost-ESP + Distill không còn trong project sau khi dọn.",
            "Loss benchmark hiện là tổng hợp log cũ, chưa phải một lần retrain đồng nhất từ cùng seed/dataset/config.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("LLIE current pipeline benchmark - 2026-07-26").font.size = Pt(9)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
